#coding=utf-8
import os
import re
import docx
import urllib
import requests

from datetime import *
from lxml import etree
from bs4 import BeautifulSoup

from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.dml import MSO_THEME_COLOR_INDEX


def endnote(document):
    p = document.add_paragraph('欢迎点赞收藏关注~')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = document.add_paragraph('https://blog.csdn.net/weixin_42815846')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_picture('比心.png', width=Inches(1))
    document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

################################################################# apartments ##################################################################
#获取页面信息
def getHtml(url):    
    # 借助user-agent
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.100 Safari/537.36'
    }
    # 借助requests获取响应对象
    response = requests.get(url=url, headers=headers)
    # 从响应对象中获取响应内容
    # 设置编码格式
    response.encoding = "utf-8"
    return response.text

rooms = []
for page in range(1, 29):
    print(f"第{page}页")
    html = getHtml(f'https://www.apartments.com/los-angeles-ca/{page}/')  
    bs4 = BeautifulSoup(html, 'html.parser')

    # 根据属性结构获取内容
    content = bs4.find_all(name="li", attrs={"class": "mortar-wrapper"})
    selecter=etree.HTML(html)
    
    for room in content:
        title = selecter.xpath('//*[@id="placardContainer"]/ul/li[1]/article/header/div[1]/a/div[1]/span')[0].xpath('string(.)')
        addr = selecter.xpath('//*[@id="placardContainer"]/ul/li[1]/article/header/div[1]/a/div[2]')[0].xpath('string(.)')
        # 房屋链接
        roomlink = selecter.xpath('//*[@id="placardContainer"]/ul/li[1]/article')[0].xpath('@data-url')[0]
        rooms.append([title, addr, roomlink])
        
        prefix = selecter.xpath('//*[@id="placardContainer"]/ul/li[1]/article/section/div')[0]
        # 图片链接
        picsrc = prefix.xpath('./div[1]/div[2]/div[2]/ul/li[1]/a')[0].xpath('@href')[0]
        # Virtual Tour
        VRtour = prefix.xpath('./div[1]/div[2]/div[2]/ul/li[2]/a')
        VRtour = VRtour[0].xpath('@href')[0] if len(VRtour)>0 else "无"
        # 免租
        freeRent =  prefix.xpath('./div[2]/div/div[1]/div')
        freeRent = freeRent[0].xpath('@data-specials-label')[0] if len(freeRent)>0 else "否"
        # 价格
        price = prefix.xpath('./div[2]/div/div[2]/div')[0].xpath('string(.)')
        # 类型
        roomtype = prefix.xpath('./div[2]/div/div[3]/div[1]')[0].xpath('string(.)')
        # 是否可用
        available = prefix.xpath('./div[2]/div/div[3]/div[2]')[0].xpath('string(.)')
        # 便利设施
        amenities = []
        for i, per1 in enumerate(prefix.xpath('./div[2]/div/div[4]/span')):
            amenities.append(per1.xpath('string(.)'))
        # 电话
        telephone = prefix.xpath('./div[2]/div/div[5]/a/span')
        telephone = telephone[0].xpath('string(.)') if len(telephone)>0 else "无"
        rooms[-1].extend([picsrc, VRtour, freeRent, price, roomtype, available, amenities, telephone])

def apartments2docx(data):
    document = Document()
    header = 'LA lease from www.apartments.com {}'.format(datetime.now().strftime('%a, %b %d %H:%M'))
    document.add_heading(header, 0)
    names = ['房屋类型', 'Availability', '价格', '是否免租', '便利设施', '室内概览', 'VR看房', '房屋链接', '联系方式']
    orders = [7, 8, 6, 5, 9, 3, 4, 2, 10]
    for i, room in enumerate(data):
        document.add_heading(f"{i+1}. "+room[0], 1)
        document.add_paragraph(room[1])
        table = document.add_table(rows=10, cols=2, style='Table Grid')

        for index, column in enumerate(table.columns):
            for cell in column.cells:
                cell.width = Inches(1) if index == 0 else Inches(5)

        table.cell(0,0).merge(table.cell(0,1))
        
        for row, obj_row in enumerate(table.rows):
            if row == 0:
                obj_row.cells[0].text = "租赁信息"
            else:
                x = orders[row-1]
                obj_row.cells[1].text = room[x] if x != 9 else '、'.join(room[x])
                obj_row.cells[0].text = names[row-1]
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        document.add_paragraph('')
    endnote(document)
    document.save('ApartmentsRooms.docx')

apartments2docx(rooms)

################################################################# tripalink ##################################################################
html = getHtml('https://tripalink.com/apartments/los-angeles/usc-off-campus-student-housing')  
bs4 = BeautifulSoup(html, 'html.parser')
# 根据属性结构获取内容
content = bs4.find_all(name="a", attrs={"target": "_blank"})

rooms = []
for room in content:
    roomlink = 'https://tripalink.com/'+room['href']
    x = room.li
    
    picsrc = x.img['src']
    addr = x.div.h2.string
    roomtype = x.div.p.span.string
    available = list(x.div.p.children)[-1].split('\n')[1].strip(' ')
    price = x.div.div.p.span.string

    rooms.append([roomlink, picsrc, addr, roomtype, available, price])

def tripalink2docx(data):
    document = Document()
    header = 'LA lease from tripalink.com {}'.format(datetime.now().strftime('%a, %b %d %H:%M'))
    document.add_heading(header, 0)
    names = ['房屋类型', 'Availability', '价格', '房屋链接', '室内概览']
    orders = [3, 4, 5, 0, 1]
    for i, room in enumerate(data):
        document.add_heading(f"{i+1}. "+room[2], 1)
        table = document.add_table(rows=6, cols=2, style='Table Grid')

        for index, column in enumerate(table.columns):
            for cell in column.cells:
                cell.width = Inches(1) if index == 0 else Inches(5)
        table.cell(0,0).merge(table.cell(0,1))
        
        for row, obj_row in enumerate(table.rows):
            if row == 0:
                obj_row.cells[0].text = "租赁信息"
            else:
                x = orders[row-1]
                obj_row.cells[1].text = room[x] if x != 9 else '、'.join(room[x])
                obj_row.cells[0].text = names[row-1]
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        document.add_paragraph('')
    endnote(document)
    document.save('TripalinkRooms.docx')

tripalink2docx(rooms)